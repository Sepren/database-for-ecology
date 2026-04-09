import pandas as pd
from docx import Document
import re
import matplotlib.pyplot as plt
from pyvis.network import Network
from collections import Counter
import glob
import sys


# === 1. ЗАГРУЗКА ТАКСОНОМИЙ ===

def find_file(pattern):
    files = glob.glob(pattern)
    return files[0] if files else None


def load_methods():
    f = find_file("*Список методов*.docx")
    if not f: return {}
    doc = Document(f)
    tax = {}
    cat = "Прочее"
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if re.match(r'^\d+\.', t):
            cat = re.sub(r'^\d+\.\s*', '', t).strip()
        elif t.startswith('-'):
            name = t.replace('-', '').split('(')[0].strip().lower()
            if name: tax[name] = cat
    return tax


def load_products():
    f = find_file("*Продукты*.docx")
    if not f: return []
    doc = Document(f)
    prods = set()
    for p in doc.paragraphs:
        t = p.text.strip()
        name = t.split(' - ')[0].strip().lower()
        if len(name) > 1: prods.add(name)
    return list(prods)


# === 2. ПОИСК СУЩНОСТЕЙ ===

def extract(text, collection):
    if not isinstance(text, str) or text.lower() == 'nan': return []
    t_low = text.lower()
    items = collection.keys() if isinstance(collection, dict) else collection
    found = [i for i in items if i in t_low]
    return list(set(found))


# === 3. ОСНОВНОЙ СКРИПТ ===

def main():
    print("--- WoodMind: Анализ данных v3 ---")

    # 1. Загрузка справочников
    m_tax = load_methods()
    p_tax = load_products()
    print(f"Словари: {len(m_tax)} методов, {len(p_tax)} продуктов.")

    # 2. Умное чтение CSV
    csv_f = find_file("*.csv")
    if not csv_f:
        print("Ошибка: CSV не найден.");
        return

    print(f"Читаем {csv_f}...")
    # Пробуем прочитать с автоматическим определением разделителя
    try:
        df = pd.read_csv(csv_f, sep=None, engine='python', encoding='utf-8-sig')
    except:
        df = pd.read_csv(csv_f, sep=None, engine='python', encoding='cp1251')

    # ИСПРАВЛЕНИЕ: Если текст ушел в индекс (как в твоем случае)
    if df['Метод'].isnull().all() or df['Метод'].astype(str).str.lower().eq('nan').all():
        print("!!! Обнаружено смещение данных. Восстанавливаем текст из индекса...")
        df['Метод'] = df.index.astype(str)

    # Очистка от пустых строк
    df = df[df['Метод'].notnull() & (df['Метод'].astype(str) != 'nan')]
    print(f"Загружено строк для анализа: {len(df)}")

    # 3. Анализ (NER)
    print("Выполняем поиск сущностей (это займет секунд 30)...")
    df['m_list'] = df['Метод'].apply(lambda x: extract(x, m_tax))
    df['p_list'] = df['Метод'].apply(lambda x: extract(x, p_tax))

    # Фильтруем только те, где нашли методы
    df_found = df[df['m_list'].apply(len) > 0].copy()
    print(f"Найдено совпадений: {len(df_found)}")

    if len(df_found) == 0:
        print("Ошибка: Совпадений не найдено. Проверь кодировку или словари.")
        return

    # 4. График ТОП-15
    all_m = [m for sub in df_found['m_list'] for m in sub]
    counts = Counter(all_m).most_common(15)

    if counts:
        labels, vals = zip(*counts)
        plt.figure(figsize=(10, 6))
        plt.barh(labels, vals, color='#6a1b9a')
        plt.title('Распределение методов (WoodMind Multi-tag Analysis)')
        plt.gca().invert_yaxis()
        plt.tight_layout()
        plt.savefig('FIG_1_TOP_METHODS.png', dpi=300)
        print("График FIG_1_TOP_METHODS.png готов.")
    else:
        print("Недостаточно данных для графика.")

    # 5. Построение Графа
    print("Генерируем интерактивный граф FIG_2_GRAPH.html...")
    net = Network(height='750px', width='100%', bgcolor='#222222', font_color='white')
    e_weights = Counter()

    for _, row in df_found.iterrows():
        for m in row['m_list']:
            net.add_node(m, label=m, color='#e67e22', group='Method')
            for p in row['p_list']:
                net.add_node(p, label=p, color='#27ae60', group='Product')
                e_weights[(m, p)] += 1

    for (m, p), w in e_weights.items():
        net.add_edge(m, p, value=w, title=f"Связей: {w}")

    net.force_atlas_2based()
    net.save_graph('FIG_2_GRAPH.html')
    print("Граф готов.")

    # 6. Сохранение результата
    df_found['m_list'] = df_found['m_list'].apply(lambda x: ', '.join(x))
    df_found['p_list'] = df_found['p_list'].apply(lambda x: ', '.join(x))
    df_found.to_csv('woodmind_final_results.csv', index=False, encoding='utf-8-sig')
    print("Результаты сохранены в woodmind_final_results.csv")


if __name__ == "__main__":
    main()